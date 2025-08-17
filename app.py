"""
PDF to Word OCR Converter - Web Application
Upload PDF files and download converted Word documents
"""
import os
import sys
import subprocess
from flask import Flask, render_template, request, send_file, jsonify, redirect, url_for, Response
from werkzeug.utils import secure_filename
import uuid
from datetime import datetime
import threading
import time

# Add the parent directory to Python path to import our converter
sys.path.append(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

try:
    import fitz  # PyMuPDF
    from docx import Document
    from spellchecker import SpellChecker
    import pytesseract
    from PIL import Image
    import io
except ImportError as e:
    print(f"Missing required package: {e}")
    print("Please install: pip install flask PyMuPDF python-docx pyspellchecker pytesseract Pillow")
    sys.exit(1)

app = Flask(__name__)
app.config['SECRET_KEY'] = 'pdf-converter-secret-key-2025'
app.config['UPLOAD_FOLDER'] = 'uploads'
app.config['DOWNLOAD_FOLDER'] = 'downloads'
app.config['MAX_CONTENT_LENGTH'] = 100 * 1024 * 1024  # 100MB max file size

# Store conversion status
conversion_status = {}

def setup_tesseract():
    """Setup Tesseract path for different environments"""
    # Check if we're in a cloud environment (Render, Heroku, etc.)
    if os.getenv('RENDER') or os.getenv('DYNO') or not os.name == 'nt':
        # On Linux/cloud platforms, tesseract should be in PATH
        print("Cloud/Linux environment detected, using system tesseract")
        try:
            # Test if tesseract is available
            result = subprocess.run(['which', 'tesseract'], capture_output=True, text=True)
            if result.returncode == 0:
                print(f"Tesseract found at: {result.stdout.strip()}")
            else:
                print("Tesseract not found in PATH")
        except:
            print("Could not check tesseract location")
        return True
    
    # Windows setup
    tesseract_paths = [
        r"C:\Program Files\Tesseract-OCR	esseract.exe",
        r"C:\Program Files (x86)\Tesseract-OCR	esseract.exe",
        r"C:\Users\{}\AppData\Local\Programs\Tesseract-OCR	esseract.exe".format(os.getenv('USERNAME'))
    ]
    
    for path in tesseract_paths:
        if os.path.exists(path):
            pytesseract.pytesseract.tesseract_cmd = path
            print(f"Tesseract found at: {path}")
            break
    else:
        print("Tesseract not found in common Windows locations")
        return False
    
    # Test Tesseract
    try:
        test_result = pytesseract.image_to_string(Image.new('RGB', (100, 100), color='white'))
        print("Tesseract test successful")
        return True
    except Exception as e:
        print(f"Tesseract test failed: {str(e)}")
        return False

def process_image_with_ocr(pil_image, image_name="unknown"):
    """Process a PIL image with OCR using multiple methods"""
    try:
        print(f"Processing {image_name} - Size: {pil_image.size}, Mode: {pil_image.mode}")
        
        # Always save debug images for troubleshooting
        debug_filename = None
        try:
            upload_folder = os.path.join(os.getcwd(), 'uploads')
            os.makedirs(upload_folder, exist_ok=True)
            debug_filename = os.path.join(upload_folder, f"debug_{image_name}.png")
            pil_image.save(debug_filename)
            print(f"Debug image saved as {debug_filename}")
        except Exception as e:
            print(f"Could not save debug image: {str(e)}")
        
        # Convert to RGB if necessary
        if pil_image.mode != 'RGB':
            pil_image = pil_image.convert('RGB')
            print(f"Converted to RGB mode")
        
        # Enhance image for better OCR
        # Resize if too small
        original_size = pil_image.size
        if pil_image.size[0] < 300 or pil_image.size[1] < 300:
            scale_factor = max(300 / pil_image.size[0], 300 / pil_image.size[1])
            new_size = (int(pil_image.size[0] * scale_factor), int(pil_image.size[1] * scale_factor))
            pil_image = pil_image.resize(new_size, Image.Resampling.LANCZOS)
            print(f"Resized image from {original_size} to: {pil_image.size}")
        
        # Try different OCR approaches
        ocr_results = []
        
        # Check if Tesseract is available first
        tesseract_available = True
        try:
            # Test Tesseract installation
            test_result = pytesseract.get_tesseract_version()
            print(f"Tesseract version: {test_result}")
            
            # Test with a simple image
            test_img = Image.new('RGB', (100, 30), color='white')
            from PIL import ImageDraw, ImageFont
            draw = ImageDraw.Draw(test_img)
            try:
                font = ImageFont.load_default()
            except:
                font = None
            
            draw.text((10, 10), "TEST", fill='black', font=font)
            test_text = pytesseract.image_to_string(test_img, config='--psm 8')
            
            if "TEST" in test_text.upper():
                print("✓ Tesseract is working correctly")
            else:
                print(f"⚠ Tesseract test failed: got '{test_text.strip()}'")
                # Try to fix common issues
                try:
                    import subprocess
                    import sys
                    
                    # Try to set PATH for Tesseract
                    import os
                    possible_paths = [
                        "/usr/bin/tesseract",
                        "/usr/local/bin/tesseract", 
                        "/opt/render/project/src/.apt/usr/bin/tesseract",
                        "/app/.apt/usr/bin/tesseract"
                    ]
                    
                    for path in possible_paths:
                        if os.path.exists(path):
                            pytesseract.pytesseract.tesseract_cmd = path
                            print(f"Set Tesseract path to: {path}")
                            break
                    
                    # Test again with correct path
                    test_text = pytesseract.image_to_string(test_img, config='--psm 8')
                    if "TEST" in test_text.upper():
                        print("✓ Tesseract working after path fix")
                    else:
                        print("✗ Tesseract still not working properly")
                        
                except Exception as path_error:
                    print(f"Could not fix Tesseract path: {path_error}")
                    
        except Exception as e:
            print(f"Tesseract not available: {str(e)}")
            tesseract_available = False
            
            # Try to auto-install/configure Tesseract for cloud environments
            try:
                print("Attempting to configure Tesseract for cloud environment...")
                
                # Check if we're on a cloud platform
                import platform
                print(f"Platform: {platform.system()} {platform.release()}")
                
                # Try common cloud paths
                cloud_paths = [
                    "/usr/bin/tesseract",
                    "/usr/local/bin/tesseract",
                    "/opt/render/project/src/.apt/usr/bin/tesseract",
                    "/app/.apt/usr/bin/tesseract",
                    "/usr/share/tesseract-ocr",
                ]
                
                for path in cloud_paths:
                    if os.path.exists(path):
                        pytesseract.pytesseract.tesseract_cmd = path
                        print(f"Found Tesseract at: {path}")
                        
                        # Test if it works now
                        try:
                            test_result = pytesseract.get_tesseract_version()
                            print(f"Tesseract now working: {test_result}")
                            tesseract_available = True
                            break
                        except:
                            continue
                
                if not tesseract_available:
                    print("Could not locate working Tesseract installation")
                    
            except Exception as config_error:
                print(f"Tesseract configuration failed: {config_error}")
        
        if tesseract_available:
            # Method 1: Try different PSM modes with various configurations
            psm_configs = [
                (6, '--psm 6 --oem 3'),
                (4, '--psm 4 --oem 3'),
                (3, '--psm 3 --oem 3'),
                (8, '--psm 8 --oem 3'),
                (11, '--psm 11 --oem 3'),
                (12, '--psm 12 --oem 3'),
                (13, '--psm 13 --oem 3'),
                (1, '--psm 1 --oem 3'),
                (6, '--psm 6 --oem 1'),
                (6, '--psm 6 --oem 2'),
                (6, '--psm 6 -c tessedit_char_whitelist=0123456789'),
                (6, '--psm 6 -c tessedit_char_whitelist=0123456789ABCDEFGHIJKLMNOPQRSTUVWXYZabcdefghijklmnopqrstuvwxyz .,():-+=[]{}'),
                (8, '--psm 8 -c tessedit_char_whitelist=0123456789ABCDEFGHIJKLMNOPQRSTUVWXYZabcdefghijklmnopqrstuvwxyz'),
                (7, '--psm 7 --oem 3'),
                (10, '--psm 10 --oem 3'),
            ]
            
            for psm, config in psm_configs:
                try:
                    ocr_text = pytesseract.image_to_string(pil_image, lang="eng", config=config)
                    cleaned_text = ocr_text.strip()
                    if cleaned_text and len(cleaned_text) > 1 and not cleaned_text.isspace():
                        ocr_results.append((f"PSM {psm}", cleaned_text))
                        print(f"OCR with {config}: '{cleaned_text[:50]}...' ({len(cleaned_text)} chars)")
                        break  # Use first successful result
                except Exception as e:
                    print(f"OCR failed with {config}: {str(e)}")
                    continue
            
            # Method 2: Try with aggressive image preprocessing if no results yet
            if not ocr_results:
                try:
                    from PIL import ImageEnhance, ImageFilter, ImageOps
                    
                    preprocessing_methods = [
                        ("Grayscale + Contrast", lambda img: ImageEnhance.Contrast(img.convert('L')).enhance(3.0)),
                        ("Grayscale + Brightness", lambda img: ImageEnhance.Brightness(img.convert('L')).enhance(1.5)),
                        ("Threshold", lambda img: img.convert('L').point(lambda x: 0 if x < 128 else 255, '1')),
                        ("Invert + Threshold", lambda img: ImageOps.invert(img.convert('L')).point(lambda x: 0 if x < 128 else 255, '1')),
                        ("Sharpen + Contrast", lambda img: ImageEnhance.Contrast(img.convert('L').filter(ImageFilter.SHARPEN)).enhance(2.5)),
                        ("Auto-contrast", lambda img: ImageOps.autocontrast(img.convert('L'))),
                    ]
                    
                    for method_name, preprocess_func in preprocessing_methods:
                        try:
                            processed_img = preprocess_func(pil_image)
                            
                            # Try multiple configs on processed image
                            for config in ['--psm 6 --oem 3', '--psm 8 --oem 3', '--psm 4 --oem 3']:
                                try:
                                    ocr_text = pytesseract.image_to_string(processed_img, lang="eng", config=config)
                                    cleaned_text = ocr_text.strip()
                                    if cleaned_text and len(cleaned_text) > 1:
                                        ocr_results.append((f"{method_name}", cleaned_text))
                                        print(f"OCR with {method_name}: '{cleaned_text[:50]}...' ({len(cleaned_text)} chars)")
                                        break
                                except Exception as e:
                                    continue
                            
                            if ocr_results:
                                break
                                
                        except Exception as e:
                            print(f"{method_name} preprocessing failed: {str(e)}")
                            continue
                    
                except Exception as e:
                    print(f"Image preprocessing failed: {str(e)}")
            
            # Method 3: Try different scale factors
            if not ocr_results:
                try:
                    scale_factors = [1.5, 2.0, 3.0, 4.0]
                    for scale in scale_factors:
                        try:
                            new_size = (int(pil_image.size[0] * scale), int(pil_image.size[1] * scale))
                            scaled_img = pil_image.resize(new_size, Image.Resampling.LANCZOS)
                            
                            # Convert to grayscale and enhance
                            gray_scaled = scaled_img.convert('L')
                            enhanced_scaled = ImageEnhance.Contrast(gray_scaled).enhance(2.0)
                            
                            config = '--psm 6 --oem 3'
                            ocr_text = pytesseract.image_to_string(enhanced_scaled, lang="eng", config=config)
                            cleaned_text = ocr_text.strip()
                            if cleaned_text and len(cleaned_text) > 1:
                                ocr_results.append((f"Scale {scale}x", cleaned_text))
                                print(f"OCR with scale {scale}x: '{cleaned_text[:50]}...' ({len(cleaned_text)} chars)")
                                break
                        except Exception as e:
                            continue
                    
                except Exception as e:
                    print(f"Scaling attempts failed: {str(e)}")
            
            # Method 4: Try different languages if still no results
            if not ocr_results:
                languages = ['eng', 'eng+deu', 'eng+fra', 'osd']
                for lang in languages:
                    try:
                        config = f'--psm 6 --oem 3 -l {lang}'
                        ocr_text = pytesseract.image_to_string(pil_image, config=config)
                        cleaned_text = ocr_text.strip()
                        if cleaned_text and len(cleaned_text) > 1:
                            ocr_results.append((f"Lang {lang}", cleaned_text))
                            print(f"OCR with {lang}: '{cleaned_text[:50]}...' ({len(cleaned_text)} chars)")
                            break
                    except Exception as e:
                        continue
        
        # Method 5: Try EasyOCR as a more reliable alternative
        if not ocr_results:
            print("Tesseract failed, trying EasyOCR...")
            try:
                try:
                    import easyocr
                    easyocr_available = True
                    print("EasyOCR available")
                except ImportError:
                    print("EasyOCR not available, installing...")
                    try:
                        import subprocess
                        import sys
                        subprocess.check_call([sys.executable, "-m", "pip", "install", "easyocr>=1.7.0", "--quiet"])
                        import easyocr
                        easyocr_available = True
                        print("EasyOCR installed successfully")
                    except Exception as install_error:
                        print(f"Could not install EasyOCR: {install_error}")
                        easyocr_available = False
                
                if easyocr_available:
                    try:
                        print("Initializing EasyOCR reader...")
                        reader = easyocr.Reader(['en'], gpu=False, verbose=False, download_enabled=True)
                        
                        # Convert PIL image to numpy array for EasyOCR
                        import numpy as np
                        img_array = np.array(pil_image)
                        
                        print("Running EasyOCR text detection...")
                        easyocr_results = reader.readtext(
                            img_array, 
                            detail=0,  # Only return text, not coordinates
                            paragraph=True,  # Group text into paragraphs
                            width_ths=0.7,  # Text width threshold
                            height_ths=0.7,  # Text height threshold
                            batch_size=1  # Reduce memory usage
                        )
                        
                        if easyocr_results:
                            combined_text = ' '.join(easyocr_results)
                            if combined_text.strip():
                                ocr_results.append(("EasyOCR", combined_text.strip()))
                                print(f"EasyOCR successful: '{combined_text[:50]}...' ({len(combined_text)} chars)")
                        else:
                            print("EasyOCR completed but found no text")
                            
                    except Exception as e:
                        print(f"EasyOCR processing failed: {str(e)}")
                else:
                    print("EasyOCR not available, skipping...")
                
            except Exception as e:
                print(f"EasyOCR error: {str(e)}")
        
        # Method 6: If all else fails, try enhanced Tesseract one more time
        if not ocr_results and tesseract_available:
            print("Trying enhanced Tesseract as last resort...")
            try:
                # Quick enhanced Tesseract configs (most likely to work)
                fast_configs = [
                    '--psm 6 --oem 3',
                    '--psm 8 --oem 3',
                    '--psm 4 --oem 3',
                    '--psm 6 --oem 1'
                ]
                
                for config in fast_configs:
                    try:
                        # Use enhanced image
                        gray_img = pil_image.convert('L')
                        from PIL import ImageEnhance
                        enhanced = ImageEnhance.Contrast(gray_img).enhance(2.5)
                        
                        ocr_text = pytesseract.image_to_string(enhanced, config=config)
                        if ocr_text.strip() and len(ocr_text.strip()) > 1:
                            ocr_results.append((f"Enhanced-Tesseract", ocr_text.strip()))
                            print(f"Enhanced Tesseract worked: '{ocr_text[:30]}...'")
                            break
                    except Exception as te:
                        continue
                        
            except Exception as e:
                print(f"Enhanced Tesseract failed: {str(e)}")
        
        # Method 7: Fallback - try to extract any visible text patterns
        if not ocr_results:
            print("All OCR methods failed, trying pattern detection...")
            try:
                # Convert to numpy array for analysis
                import numpy as np
                img_array = np.array(pil_image)
                
                # Check if image has text-like patterns (high contrast areas)
                if len(img_array.shape) == 3:
                    gray_array = np.dot(img_array[...,:3], [0.2989, 0.5870, 0.1140])
                else:
                    gray_array = img_array
                
                # Calculate contrast and other metrics
                contrast = gray_array.std()
                mean_brightness = gray_array.mean()
                
                # Try one more aggressive attempt with extreme preprocessing
                try:
                    from PIL import ImageOps
                    
                    # Convert to pure black and white
                    gray_pil = pil_image.convert('L')
                    
                    # Try different threshold values
                    thresholds = [100, 128, 150, 180, 200]
                    for threshold in thresholds:
                        try:
                            # Apply threshold
                            bw_img = gray_pil.point(lambda x: 0 if x < threshold else 255, '1')
                            
                            # Try OCR on binary image
                            config = '--psm 6 --oem 3 -c tessedit_char_whitelist=0123456789ABCDEFGHIJKLMNOPQRSTUVWXYZabcdefghijklmnopqrstuvwxyz'
                            ocr_text = pytesseract.image_to_string(bw_img, config=config)
                            cleaned_text = ocr_text.strip()
                            
                            if cleaned_text and len(cleaned_text) > 1:
                                ocr_results.append((f"Threshold {threshold}", cleaned_text))
                                print(f"OCR with threshold {threshold}: '{cleaned_text[:50]}...' ({len(cleaned_text)} chars)")
                                break
                        except Exception as e:
                            continue
                    
                    # If still no results, try inverted image
                    if not ocr_results:
                        try:
                            inverted_img = ImageOps.invert(gray_pil)
                            config = '--psm 6 --oem 3'
                            ocr_text = pytesseract.image_to_string(inverted_img, config=config)
                            cleaned_text = ocr_text.strip()
                            
                            if cleaned_text and len(cleaned_text) > 1:
                                ocr_results.append(("Inverted", cleaned_text))
                                print(f"OCR with inverted image: '{cleaned_text[:50]}...' ({len(cleaned_text)} chars)")
                        except Exception as e:
                            print(f"Inverted OCR failed: {str(e)}")
                    
                except Exception as e:
                    print(f"Aggressive preprocessing failed: {str(e)}")
                
                # If still no results, provide detailed analysis
                if not ocr_results:
                    print(f"Image analysis - Contrast: {contrast:.1f}, Brightness: {mean_brightness:.1f}")
                    
                    # Try to detect text regions using simple edge detection
                    try:
                        from PIL import ImageFilter
                        edges = gray_pil.filter(ImageFilter.FIND_EDGES)
                        edge_array = np.array(edges)
                        edge_density = (edge_array > 0).sum() / edge_array.size
                        print(f"Edge density: {edge_density:.3f}")
                        
                        # Try to save processed images for manual inspection
                        try:
                            processed_folder = os.path.join(os.getcwd(), 'uploads', 'processed')
                            os.makedirs(processed_folder, exist_ok=True)
                            
                            # Save different processed versions
                            gray_pil.save(os.path.join(processed_folder, f"gray_{image_name}.png"))
                            edges.save(os.path.join(processed_folder, f"edges_{image_name}.png"))
                            
                            # Try binary threshold
                            binary = gray_pil.point(lambda x: 0 if x < 128 else 255, '1')
                            binary.save(os.path.join(processed_folder, f"binary_{image_name}.png"))
                            
                            print(f"Processed images saved to {processed_folder}")
                            
                        except Exception as e:
                            print(f"Could not save processed images: {str(e)}")
                            
                    except Exception as e:
                        print(f"Edge detection failed: {str(e)}")
                    
                    if contrast > 30:  # Likely has text
                        analysis_text = f"Image contains text-like patterns (Contrast: {contrast:.1f}, Brightness: {mean_brightness:.1f}) but OCR extraction failed.\n"
                        analysis_text += f"This could indicate:\n"
                        analysis_text += f"• Mathematical symbols or special characters\n"
                        analysis_text += f"• Non-standard fonts or handwriting\n"
                        analysis_text += f"• Complex layouts or formatting\n"
                        analysis_text += f"• Image quality issues\n"
                        analysis_text += f"• OCR engine limitations\n"
                        analysis_text += f"Image size: {pil_image.size[0]}x{pil_image.size[1]} pixels"
                        
                        if debug_filename:
                            analysis_text += f"\nDebug image saved for manual inspection"
                            
                        ocr_results.append(("Analysis", analysis_text))
                    else:
                        ocr_results.append(("Analysis", f"Low contrast image (Contrast: {contrast:.1f}) - may not contain readable text"))
                    
            except Exception as e:
                print(f"Pattern analysis failed: {str(e)}")
                ocr_results.append(("Fallback", f"Image detected but text extraction failed - Error: {str(e)}"))
        
        # Return the best result
        if ocr_results:
            # Prioritize actual text over analysis messages
            text_results = [r for r in ocr_results if not r[1].startswith('[') and not r[1].startswith('Image')]
            if text_results:
                best_method, best_text = max(text_results, key=lambda x: len(x[1]))
            else:
                best_method, best_text = ocr_results[0]  # Use first result (likely analysis)
            
            print(f"Best OCR result for {image_name} using {best_method}: {len(best_text)} characters")
            return best_text
        else:
            print(f"No text found in {image_name}")
            return f"[Image detected but no text could be extracted from {image_name}]"
            
    except Exception as e:
        print(f"Error processing {image_name}: {str(e)}")
        return f"[Error processing image: {str(e)}]"

def extract_text_from_page(page):
    """Extract text from both regular text and images in a PDF page"""
    regular_text = page.get_text().strip()
    image_text = ""
    
    # First, try to get images from the page using multiple methods
    try:
        # Method 1: Get image list (most common)
        image_list = page.get_images(full=True)
        print(f"Method 1 - Found {len(image_list)} images in image list")
        
        # Method 2: Check for drawing objects that might contain images
        drawings = page.get_drawings()
        print(f"Method 2 - Found {len(drawings)} drawing objects")
        
        # Method 3: Check page dictionary for image blocks
        page_dict = page.get_text("dict")
        image_blocks = []
        for block in page_dict.get("blocks", []):
            if "type" in block and block["type"] == 1:  # Image block
                image_blocks.append(block)
        print(f"Method 3 - Found {len(image_blocks)} image blocks in page dict")
        
        # Method 4: Check page contents for image references
        try:
            page_obj = page.get_contents()
            if page_obj:
                content_stream = page_obj[0].get_buffer()
                if b"/Im" in content_stream or b"/Image" in content_stream or b"Do" in content_stream:
                    print("Method 4 - Found image references in page contents")
        except:
            pass
        
        # If no images found with standard method, try alternative approach
        if not image_list:
            print("No images found with standard method, trying alternative extraction...")
            # Try to render the page and check if it contains non-text content
            try:
                # Get page as pixmap to check for visual content
                mat = fitz.Matrix(1.0, 1.0)
                pix = page.get_pixmap(matrix=mat)
                
                # If page has visual content but no extractable text, it might be all images
                if pix.size > 0 and len(regular_text) < 50:
                    print("Page appears to contain mainly visual content, treating as image")
                    # Render page at higher resolution for OCR
                    mat = fitz.Matrix(2.0, 2.0)
                    pix = page.get_pixmap(matrix=mat)
                    img_data = pix.tobytes("png")
                    pil_image = Image.open(io.BytesIO(img_data))
                    
                    # Process as full page image
                    ocr_text = process_image_with_ocr(pil_image, "full_page")
                    if ocr_text and ocr_text.strip():
                        image_text += f"\n[Text from full page image:]\n{ocr_text.strip()}\n"
                        print(f"Extracted text from full page: {len(ocr_text)} characters")
                    
                    pix = None
            except Exception as e:
                print(f"Alternative extraction failed: {str(e)}")
        
        # Process individual images from image_list
        for img_index, img in enumerate(image_list):
            try:
                xref = img[0]
                print(f"Processing image {img_index + 1} with xref {xref}")
                
                # Try to extract the image using different methods
                try:
                    # Method A: extract_image (newer PyMuPDF)
                    base_image = page.parent.extract_image(xref)
                    image_bytes = base_image["image"]
                    image_ext = base_image["ext"]
                    print(f"Image {img_index + 1}: format={image_ext}, size={len(image_bytes)} bytes")
                    
                    # Convert to PIL Image
                    pil_image = Image.open(io.BytesIO(image_bytes))
                    
                except:
                    # Method B: Pixmap extraction (fallback)
                    print("Using pixmap fallback for image extraction")
                    pix = fitz.Pixmap(page.parent, xref)
                    
                    if pix.n - pix.alpha < 4:
                        img_data = pix.tobytes("png")
                        pil_image = Image.open(io.BytesIO(img_data))
                    else:
                        # Convert CMYK to RGB
                        pix = fitz.Pixmap(fitz.csRGB, pix)
                        img_data = pix.tobytes("png")
                        pil_image = Image.open(io.BytesIO(img_data))
                    
                    pix = None
                
                print(f"PIL Image - Size: {pil_image.size}, Mode: {pil_image.mode}")
                
                # Process the image with OCR
                ocr_text = process_image_with_ocr(pil_image, f"image_{img_index + 1}")
                
                if ocr_text and ocr_text.strip():
                    image_text += f"\n[Text from image {img_index + 1}:]\n{ocr_text.strip()}\n"
                    print(f"Successfully extracted text from image {img_index + 1}: {len(ocr_text)} characters")
                else:
                    print(f"No text found in image {img_index + 1}")
                    
            except Exception as img_error:
                print(f"Error processing image {img_index + 1}: {str(img_error)}")
                continue
            
    except Exception as e:
        print(f"Error accessing images: {str(e)}")
    
    # Fallback: Try full page OCR if no regular text and no image text
    if not regular_text and not image_text:
        try:
            print("Attempting full page OCR as last resort...")
            # Render the entire page as an image
            mat = fitz.Matrix(2.0, 2.0)  # Higher resolution
            pix = page.get_pixmap(matrix=mat)
            img_data = pix.tobytes("png")
            pil_image = Image.open(io.BytesIO(img_data))
            
            # Try different configurations for full page
            for psm in [1, 3, 4, 6]:
                try:
                    config = f'--psm {psm}'
                    page_ocr_text = pytesseract.image_to_string(pil_image, lang="eng", config=config)
                    if page_ocr_text.strip() and len(page_ocr_text.strip()) > 10:
                        print(f"Full page OCR successful with PSM {psm}: {len(page_ocr_text)} characters")
                        return page_ocr_text.strip(), f"Full page OCR (PSM {psm})", True, True
                except Exception as e:
                    continue
            
            pix = None
                
        except Exception as ocr_error:
            print(f"Full page OCR failed: {str(ocr_error)}")
    
    # Count images found
    total_images = len(image_list)
    if total_images == 0 and len(regular_text) < 50:
        # Might be a visual-heavy page, count as having images
        total_images = 1
    
    # Combine results
    combined_text = regular_text
    if image_text:
        combined_text = regular_text + "\n" + image_text
    
    extraction_method = "Text"
    has_images = total_images > 0
    used_ocr = False
    
    if image_text and regular_text:
        extraction_method = "Text + Image OCR"
        used_ocr = True
    elif image_text and not regular_text:
        extraction_method = "Image OCR only"
        used_ocr = True
    elif not regular_text and not image_text and total_images > 0:
        extraction_method = "Images detected (OCR failed)"
    elif not regular_text and not image_text:
        extraction_method = "No text found"
    
    # Consider OCR used if we found images and tried to process them
    if total_images > 0 and "[Image" in combined_text:
        used_ocr = True
    
    return combined_text.strip(), extraction_method, has_images, used_ocr

def test_ocr_functionality():
    """Test if OCR is working properly"""
    print("=" * 50)
    print("COMPREHENSIVE OCR FUNCTIONALITY TEST")
    print("=" * 50)
    
    try:
        # First, check Tesseract availability
        print("\n1. Testing Tesseract Installation:")
        try:
            version = pytesseract.get_tesseract_version()
            print(f"   ✓ Tesseract version: {version}")
            tesseract_ok = True
        except Exception as e:
            print(f"   ✗ Tesseract not available: {str(e)}")
            tesseract_ok = False
            
            # Try to find Tesseract in common locations
            import os
            common_paths = [
                "/usr/bin/tesseract",
                "/usr/local/bin/tesseract", 
                "/opt/render/project/src/.apt/usr/bin/tesseract",
                "/app/.apt/usr/bin/tesseract"
            ]
            
            for path in common_paths:
                if os.path.exists(path):
                    print(f"   Found Tesseract at: {path}")
                    pytesseract.pytesseract.tesseract_cmd = path
                    try:
                        version = pytesseract.get_tesseract_version()
                        print(f"   ✓ Tesseract working with path: {version}")
                        tesseract_ok = True
                        break
                    except:
                        continue
        
        # Test environment
        print("\n2. Environment Information:")
        import platform, sys
        print(f"   Platform: {platform.system()} {platform.release()}")
        print(f"   Python: {sys.version}")
        print(f"   Working directory: {os.getcwd()}")
        
        # Check if we can write files
        try:
            test_file = "test_write.txt"
            with open(test_file, 'w') as f:
                f.write("test")
            os.remove(test_file)
            print("   ✓ File system is writable")
        except Exception as e:
            print(f"   ⚠ File system issue: {e}")
        
        # Create a simple test image with text
        print("\n3. Creating Test Image:")
        from PIL import Image, ImageDraw, ImageFont
        
        # Create a white image with text
        img = Image.new('RGB', (400, 150), color='white')
        draw = ImageDraw.Draw(img)
        
        # Try to use a font, fall back to default if not available
        try:
            font = ImageFont.truetype("arial.ttf", 32)
            print("   ✓ Using TrueType font")
        except:
            try:
                font = ImageFont.load_default()
                print("   ✓ Using default font")
            except:
                font = None
                print("   ⚠ No font available")
        
        draw.text((20, 50), "Test OCR Text 123", fill='black', font=font)
        
        # Save test image for debugging
        try:
            img.save("test_ocr_image.png")
            print("   ✓ Test image saved as test_ocr_image.png")
        except Exception as e:
            print(f"   ⚠ Could not save test image: {e}")
        
        # Test OCR if Tesseract is available
        if tesseract_ok:
            print("\n4. Testing Tesseract OCR:")
            test_results = []
            
            # Method 1: Different PSM modes
            for psm in [6, 8, 7, 13]:
                try:
                    config = f'--psm {psm} --oem 3'
                    test_text = pytesseract.image_to_string(img, lang="eng", config=config)
                    cleaned = test_text.strip()
                    if cleaned:
                        test_results.append(f"PSM {psm}: '{cleaned}'")
                        print(f"   ✓ PSM {psm}: '{cleaned}'")
                    else:
                        print(f"   ✗ PSM {psm}: No text detected")
                except Exception as e:
                    print(f"   ✗ PSM {psm}: Failed - {str(e)}")
            
            # Method 2: Simple OCR
            try:
                simple_text = pytesseract.image_to_string(img)
                cleaned = simple_text.strip()
                if cleaned:
                    test_results.append(f"Simple: '{cleaned}'")
                    print(f"   ✓ Simple OCR: '{cleaned}'")
                else:
                    print("   ✗ Simple OCR: No text detected")
            except Exception as e:
                print(f"   ✗ Simple OCR: Failed - {str(e)}")
            
            # Check if any method worked
            success_count = sum(1 for result in test_results if any(word in result.upper() for word in ["TEST", "OCR", "TEXT", "123"]))
            
            if success_count > 0:
                print(f"\n   ✓ Tesseract OCR working! ({success_count} successful methods)")
                tesseract_success = True
            else:
                print("\n   ✗ Tesseract OCR not working - no text detected correctly")
                tesseract_success = False
        else:
            tesseract_success = False
        
        # Test EasyOCR
        print("\n5. Testing EasyOCR:")
        try:
            import easyocr
            print("   ✓ EasyOCR available")
            
            try:
                reader = easyocr.Reader(['en'], gpu=False, verbose=False, download_enabled=True)
                print("   ✓ EasyOCR reader initialized")
                
                import numpy as np
                img_array = np.array(img)
                easyocr_results = reader.readtext(img_array, detail=0, paragraph=True)
                
                if easyocr_results:
                    combined_text = ' '.join(easyocr_results)
                    print(f"   ✓ EasyOCR result: '{combined_text}'")
                    easyocr_success = True
                else:
                    print("   ✗ EasyOCR: No text detected")
                    easyocr_success = False
                    
            except Exception as e:
                print(f"   ✗ EasyOCR processing failed: {str(e)}")
                easyocr_success = False
                
        except ImportError:
            print("   ⚠ EasyOCR not available")
            easyocr_success = False
        except Exception as e:
            print(f"   ✗ EasyOCR error: {str(e)}")
            easyocr_success = False
        
        # Final assessment
        print("\n" + "=" * 50)
        print("FINAL ASSESSMENT:")
        print("=" * 50)
        
        if tesseract_success or easyocr_success:
            print("✓ OCR functionality is working!")
            if tesseract_success:
                print("  - Tesseract: Working")
            if easyocr_success:
                print("  - EasyOCR: Working")
            return True
        else:
            print("✗ OCR functionality is NOT working!")
            print("  - Tesseract: Failed" if tesseract_ok else "  - Tesseract: Not available")
            print("  - EasyOCR: Failed" if not easyocr_success else "  - EasyOCR: Not tested")
            return False
            
    except Exception as e:
        print(f"✗ OCR test completely failed: {str(e)}")
        import traceback
        traceback.print_exc()
        return False

def extract_text_from_page_with_timeout(page, timeout_seconds=60):
    """Extract text from page with timeout"""
    import threading
    import time
    
    result = [None]  # Use list to store result from thread
    exception = [None]  # Store any exception
    
    def extraction_worker():
        try:
            result[0] = extract_text_from_page(page)
        except Exception as e:
            exception[0] = e
    
    # Start extraction in a separate thread
    thread = threading.Thread(target=extraction_worker)
    thread.daemon = True
    thread.start()
    
    # Wait for completion or timeout
    thread.join(timeout_seconds)
    
    if thread.is_alive():
        # Thread is still running, extraction timed out
        print(f"Text extraction timed out after {timeout_seconds} seconds")
        raise TimeoutError(f"Text extraction timed out after {timeout_seconds} seconds")
    
    if exception[0]:
        raise exception[0]
    
    if result[0] is None:
        raise Exception("Text extraction failed for unknown reason")
    
    return result[0]

def convert_pdf_to_word_web(pdf_path, output_path, job_id):
    """Convert PDF to Word with progress tracking for web interface"""
    import time
    start_time = time.time()
    max_processing_time = 300  # 5 minutes max
    
    try:
        conversion_status[job_id] = {
            'status': 'starting',
            'progress': 0,
            'message': 'Initializing conversion...',
            'current_page': 0,
            'total_pages': 0,
            'start_time': start_time
        }
        
        # Setup Tesseract
        setup_tesseract()
        
        # Check if we're taking too long
        if time.time() - start_time > max_processing_time:
            raise TimeoutError("Conversion timeout - taking too long")
        
        # Open PDF
        pdf_document = fitz.open(pdf_path)
        total_pages = len(pdf_document)
        
        conversion_status[job_id].update({
            'total_pages': total_pages,
            'message': f'Processing {total_pages} pages...'
        })
        
        # Create Word document
        doc = Document()
        pdf_name = os.path.basename(pdf_path)
        doc.add_heading(f'{pdf_name}', 0)
        
        total_words = 0
        pages_with_text = 0
        pages_with_images = 0
        pages_with_ocr = 0
        
        for page_num in range(total_pages):
            # Check timeout again
            elapsed_time = time.time() - start_time
            if elapsed_time > max_processing_time:
                conversion_status[job_id].update({
                    'status': 'error',
                    'message': f'Conversion timeout after {elapsed_time:.1f} seconds'
                })
                raise TimeoutError(f"Conversion timeout after {elapsed_time:.1f} seconds")
            
            # Update progress
            progress = int((page_num / total_pages) * 100)
            conversion_status[job_id].update({
                'status': 'processing',
                'progress': progress,
                'current_page': page_num + 1,
                'message': f'Processing page {page_num + 1} of {total_pages}... ({elapsed_time:.1f}s elapsed)'
            })
            
            # Extract text with timeout for individual pages
            page = pdf_document[page_num]
            
            try:
                # Set a per-page timeout
                page_start = time.time()
                page_timeout = 60  # 1 minute per page max
                
                text, extraction_method, has_images, used_ocr = extract_text_from_page_with_timeout(page, page_timeout)
                
                page_elapsed = time.time() - page_start
                print(f"Page {page_num + 1} processed in {page_elapsed:.1f}s using {extraction_method}")
                
            except TimeoutError as te:
                print(f"Page {page_num + 1} timed out: {str(te)}")
                text = f"[Page {page_num + 1} processing timed out - may contain complex content]"
                extraction_method = "Timeout"
                has_images = False
                used_ocr = False
            except Exception as pe:
                print(f"Page {page_num + 1} error: {str(pe)}")
                text = f"[Page {page_num + 1} processing failed: {str(pe)}]"
                extraction_method = "Error"
                has_images = False
                used_ocr = False
            
            # Track extraction methods
            if used_ocr:
                pages_with_ocr += 1
            if has_images:
                pages_with_images += 1
            
            if text:
                pages_with_text += 1
                word_count = len(text.split())
                total_words += word_count
                
                # Add to document
                doc.add_heading(f'Page {page_num + 1} ({extraction_method})', level=1)
                doc.add_paragraph(text)
                doc.add_page_break()
            else:
                # Add info about empty page
                doc.add_heading(f'Page {page_num + 1} (No text found)', level=1)
                doc.add_paragraph("[This page appears to be blank or contains no extractable text]")
                doc.add_page_break()
        
        # Add summary
        doc.add_heading('Document Information', level=1)
        doc.add_paragraph(f"Source File: {pdf_name}")
        doc.add_paragraph(f"Total Pages: {total_pages}")
        doc.add_paragraph(f"Pages with Text: {pages_with_text}")
        doc.add_paragraph(f"Pages with Images: {pages_with_images}")
        doc.add_paragraph(f"Pages using OCR: {pages_with_ocr}")
        doc.add_paragraph(f"Total Words: {total_words}")
        doc.add_paragraph(f"Converted: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        
        if pages_with_ocr == 0 and pages_with_images > 0:
            doc.add_paragraph("")
            doc.add_paragraph("⚠️ Note: Some images were detected but OCR may not have worked properly.")
            doc.add_paragraph("This could be due to:")
            doc.add_paragraph("• Tesseract OCR not being properly installed")
            doc.add_paragraph("• Images containing non-text content")
            doc.add_paragraph("• Poor image quality or resolution")
            doc.add_paragraph("• Unsupported image formats")
        
        # Save document
        doc.save(output_path)
        pdf_document.close()
        
        # Update final status
        conversion_status[job_id].update({
            'status': 'completed',
            'progress': 100,
            'message': 'Conversion completed successfully!',
            'output_file': os.path.basename(output_path),
            'stats': {
                'total_pages': total_pages,
                'pages_with_text': pages_with_text,
                'total_words': total_words
            }
        })
        
        return True
        
    except Exception as e:
        conversion_status[job_id].update({
            'status': 'error',
            'progress': 0,
            'message': f'Error: {str(e)}'
        })
        return False

@app.route('/')
def index():
    """Main page with upload form"""
    return render_template('index.html')

@app.route('/upload', methods=['POST'])
def upload_file():
    """Handle file upload and start conversion"""
    if 'pdf_file' not in request.files:
        return jsonify({'error': 'No file selected'}), 400
    
    file = request.files['pdf_file']
    if file.filename == '':
        return jsonify({'error': 'No file selected'}), 400
    
    if file and file.filename.lower().endswith('.pdf'):
        # Generate unique job ID
        job_id = str(uuid.uuid4())
        
        # Save uploaded file
        filename = secure_filename(file.filename)
        pdf_path = os.path.join(app.config['UPLOAD_FOLDER'], f"{job_id}_{filename}")
        file.save(pdf_path)
        
        # Generate output filename
        output_filename = f"{os.path.splitext(filename)[0]}_converted.docx"
        output_path = os.path.join(app.config['DOWNLOAD_FOLDER'], f"{job_id}_{output_filename}")
        
        # Start conversion in background thread
        thread = threading.Thread(
            target=convert_pdf_to_word_web,
            args=(pdf_path, output_path, job_id)
        )
        thread.start()
        
        return jsonify({
            'job_id': job_id,
            'message': 'Upload successful! Conversion started.',
            'filename': filename
        })
    
    return jsonify({'error': 'Please select a valid PDF file'}), 400

@app.route('/progress/<job_id>')
def get_progress(job_id):
    """Get conversion progress"""
    if job_id in conversion_status:
        return jsonify(conversion_status[job_id])
    return jsonify({'error': 'Job not found'}), 404

@app.route('/download/<job_id>')
def download_file(job_id):
    """Download converted Word document"""
    print(f"Download request for job_id: {job_id}")
    
    if job_id in conversion_status:
        status = conversion_status[job_id]
        print(f"Job status: {status}")
        
        if status['status'] == 'completed':
            output_file = status['output_file']
            # The output_file already has the job_id prefix, so use it directly
            file_path = os.path.join(app.config['DOWNLOAD_FOLDER'], output_file)
            
            print(f"Looking for file: {file_path}")
            print(f"File exists: {os.path.exists(file_path)}")
            
            if os.path.exists(file_path):
                # Create a clean filename for download (remove job_id prefix)
                clean_filename = output_file.replace(f"{job_id}_", "")
                print(f"Sending file as: {clean_filename}")
                
                # Simple file response
                def generate():
                    with open(file_path, 'rb') as f:
                        while True:
                            data = f.read(4096)
                            if not data:
                                break
                            yield data
                
                return Response(
                    generate(),
                    mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                    headers={
                        'Content-Disposition': f'attachment; filename="{clean_filename}"',
                        'Content-Type': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
                    }
                )
            else:
                print("File not found on disk")
                return jsonify({'error': f'File not found: {output_file}'}), 404
        else:
            print(f"Job not completed. Status: {status['status']}")
            return jsonify({'error': f'Conversion not completed. Status: {status["status"]}'}), 400
    else:
        print("Job ID not found in conversion_status")
        return jsonify({'error': 'Job not found'}), 404

@app.route('/test')
def test_page():
    """Test page for debugging downloads"""
    return render_template('test.html')

@app.route('/test-download')
def test_download():
    """Test download functionality with existing file"""
    download_folder = app.config['DOWNLOAD_FOLDER']
    files = os.listdir(download_folder) if os.path.exists(download_folder) else []
    
    if files:
        # Use the first available file for testing
        test_file = files[0]
        file_path = os.path.join(download_folder, test_file)
        
        # Extract original filename from the prefixed filename
        original_name = test_file.split('_', 1)[1] if '_' in test_file else test_file
        
        try:
            return send_file(
                os.path.abspath(file_path),
                as_attachment=True,
                attachment_filename=original_name
            )
        except Exception as e:
            return jsonify({'error': f'Download failed: {str(e)}'}), 500
    else:
        return jsonify({'error': 'No files available for download'}), 404

@app.route('/debug/<job_id>')
def debug_job(job_id):
    """Debug endpoint to check job status and files"""
    info = {
        'job_id': job_id,
        'job_exists': job_id in conversion_status,
        'download_folder': app.config['DOWNLOAD_FOLDER'],
        'download_folder_exists': os.path.exists(app.config['DOWNLOAD_FOLDER']),
        'files_in_download_folder': []
    }
    
    if os.path.exists(app.config['DOWNLOAD_FOLDER']):
        info['files_in_download_folder'] = os.listdir(app.config['DOWNLOAD_FOLDER'])
    
    if job_id in conversion_status:
        info['job_status'] = conversion_status[job_id]
    
    return jsonify(info)

@app.route('/test-ocr')
def test_ocr_endpoint():
    """Test OCR functionality and return detailed results"""
    print("OCR test endpoint called")
    
    # Capture all print output
    import io
    import contextlib
    
    f = io.StringIO()
    with contextlib.redirect_stdout(f):
        ocr_working = test_ocr_functionality()
    
    output = f.getvalue()
    
    return f"""
    <html>
    <head><title>OCR Test Results</title></head>
    <body>
        <h1>OCR Test Results</h1>
        <h2>Status: {'✓ WORKING' if ocr_working else '✗ NOT WORKING'}</h2>
        <h3>Detailed Output:</h3>
        <pre style="background: #f0f0f0; padding: 10px; white-space: pre-wrap;">{output}</pre>
        <p><a href="/">Back to Home</a></p>
    </body>
    </html>
    """

@app.route('/cleanup')
def cleanup_files():
    """Clean up old files (optional endpoint)"""
    try:
        # Remove files older than 1 hour
        current_time = time.time()
        
        for folder in [app.config['UPLOAD_FOLDER'], app.config['DOWNLOAD_FOLDER']]:
            for filename in os.listdir(folder):
                file_path = os.path.join(folder, filename)
                if os.path.isfile(file_path):
                    file_age = current_time - os.path.getctime(file_path)
                    if file_age > 3600:  # 1 hour
                        os.remove(file_path)
        
        return jsonify({'message': 'Cleanup completed'})
    except Exception as e:
        return jsonify({'error': str(e)}), 500

def diagnose_tesseract():
    """Diagnose Tesseract installation and capabilities"""
    print("\n" + "="*50)
    print("TESSERACT DIAGNOSTIC REPORT")
    print("="*50)
    
    try:
        # Check if pytesseract is available
        import pytesseract
        print("✓ pytesseract module imported successfully")
        
        # Check Tesseract version
        try:
            version = pytesseract.get_tesseract_version()
            print(f"✓ Tesseract version: {version}")
        except Exception as e:
            print(f"✗ Cannot get Tesseract version: {str(e)}")
        
        # Check available languages
        try:
            langs = pytesseract.get_languages()
            print(f"✓ Available languages: {langs}")
        except Exception as e:
            print(f"✗ Cannot get languages: {str(e)}")
        
        # Test basic OCR
        try:
            from PIL import Image, ImageDraw, ImageFont
            
            # Create test image
            img = Image.new('RGB', (400, 100), color='white')
            draw = ImageDraw.Draw(img)
            
            try:
                font = ImageFont.truetype("arial.ttf", 20)
            except:
                font = ImageFont.load_default()
            
            draw.text((10, 30), "Test 123 ABC", fill='black', font=font)
            
            # Test different OCR methods
            methods = [
                ("Simple", lambda: pytesseract.image_to_string(img)),
                ("PSM 6", lambda: pytesseract.image_to_string(img, config='--psm 6')),
                ("PSM 8", lambda: pytesseract.image_to_string(img, config='--psm 8')),
                ("Numbers only", lambda: pytesseract.image_to_string(img, config='--psm 6 -c tessedit_char_whitelist=0123456789')),
            ]
            
            for method_name, method_func in methods:
                try:
                    result = method_func()
                    if result.strip():
                        print(f"✓ {method_name}: '{result.strip()}'")
                    else:
                        print(f"✗ {method_name}: No text detected")
                except Exception as e:
                    print(f"✗ {method_name}: Failed - {str(e)}")
        
        except Exception as e:
            print(f"✗ OCR test failed: {str(e)}")
            
    except Exception as e:
        print(f"✗ pytesseract not available: {str(e)}")
    
    print("="*50)
    print()

if __name__ == '__main__':
    # Create directories if they don't exist
    os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)
    os.makedirs(app.config['DOWNLOAD_FOLDER'], exist_ok=True)
    
    # Test OCR functionality
    print("Testing OCR setup...")
    setup_tesseract()
    diagnose_tesseract()
    test_ocr_functionality()
    
    port = int(os.environ.get('PORT', 5000))
    debug = os.environ.get('FLASK_ENV') != 'production'
    
    print("🌐 PDF to Word OCR Converter Web App")
    print("=" * 50)
    print("Starting web server...")
    if debug:
        print("Open your browser and go to: http://localhost:5000")
    else:
        print("Production mode - serving on all interfaces")
    print("=" * 50)
    
    app.run(debug=debug, host='0.0.0.0', port=port)
